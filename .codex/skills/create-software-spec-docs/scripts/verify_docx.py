#!/usr/bin/env python3
"""Verify DOCX/template package integrity and Markdown-to-Word structure parity."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from pathlib import Path
from zipfile import ZipFile

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover - environment diagnostic
    raise SystemExit("python-docx is required. Use the bundled workspace Python or install python-docx.") from exc


def markdown_headings(text: str) -> list[tuple[str, str]]:
    result = []
    for line in text.splitlines():
        match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if match:
            result.append((f"Heading {len(match.group(1)) - 1}", match.group(2).strip()))
    return result


def markdown_tables(text: str) -> list[tuple[str, ...]]:
    lines = text.splitlines()
    result = []
    for index, line in enumerate(lines[:-1]):
        if not line.startswith("|") or not lines[index + 1].startswith("|"):
            continue
        separators = [cell.strip() for cell in lines[index + 1].strip("|").split("|")]
        if separators and all(re.fullmatch(r":?-{3,}:?", cell) for cell in separators):
            result.append(tuple(cell.strip().replace("`", "") for cell in line.strip("|").split("|")))
    return result


def inventory(path: Path) -> dict[str, str]:
    with ZipFile(path) as package:
        return {name: hashlib.sha256(package.read(name)).hexdigest() for name in package.namelist()}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--markdown", type=Path, required=True)
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--json", type=Path)
    args = parser.parse_args()

    markdown = args.markdown.read_text(encoding="utf-8")
    doc = Document(args.docx)
    template = Document(args.template)
    findings = []

    word_headings = [(p.style.name, p.text) for p in doc.paragraphs if p.style.name.startswith("Heading")]
    toc_heading = next((p.text for p in template.paragraphs if p.style.name == "Heading 1"), None)
    expected_headings = [("Heading 1", toc_heading)] + markdown_headings(markdown)
    if word_headings != expected_headings:
        findings.append({"code": "heading-parity", "message": "DOCX 标题层级或顺序与 Markdown 不一致"})

    word_tables = [tuple(cell.text.strip().replace("`", "") for cell in table.rows[0].cells) for table in doc.tables]
    expected_tables = markdown_tables(markdown)
    if len(word_tables) < 1 or word_tables[1:] != expected_tables:
        findings.append({"code": "table-parity", "message": "DOCX 正文表头或顺序与 Markdown 不一致"})
    if not doc.tables or len(doc.tables[0].rows) != 5 or len(doc.tables[0].columns) != 2:
        findings.append({"code": "cover-table", "message": "DOCX 未保留模板 5x2 封面信息表"})

    with ZipFile(args.docx) as package:
        document_xml = package.read("word/document.xml")
        if b"TOC" not in document_xml:
            findings.append({"code": "toc-field", "message": "DOCX 未保留目录域"})

    base, final = inventory(args.template), inventory(args.docx)
    changed = sorted(name for name in base if base[name] != final.get(name))
    unexpected = sorted(set(changed) - {"word/document.xml", "word/header1.xml"})
    if set(base) != set(final) or unexpected:
        findings.append(
            {
                "code": "package-integrity",
                "message": "DOCX 修改了模板正文和页眉之外的包组件",
                "changed": changed,
            }
        )

    full_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    placeholders = sorted(set(re.findall(r"〈[^〉]+〉|(?:FR|BD|TC)-XXX(?:-\d{3})?", full_text)))
    if placeholders:
        findings.append({"code": "template-placeholder", "message": "DOCX 仍包含模板占位符", "values": placeholders})

    report = {
        "passed": not findings,
        "markdown": str(args.markdown.resolve()),
        "docx": str(args.docx.resolve()),
        "template": str(args.template.resolve()),
        "heading_count": len(word_headings),
        "body_table_count": max(0, len(word_tables) - 1),
        "changed_parts": changed,
        "findings": findings,
    }
    if args.json:
        args.json.parent.mkdir(parents=True, exist_ok=True)
        args.json.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["passed"] else 1


if __name__ == "__main__":
    sys.exit(main())
