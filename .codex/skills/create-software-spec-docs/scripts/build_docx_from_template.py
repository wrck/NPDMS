#!/usr/bin/env python3
"""Build a DOCX from filled Markdown while preserving the approved DOCX template package."""

from __future__ import annotations

import argparse
import hashlib
import os
import re
import shutil
import sys
import tempfile
from pathlib import Path
from zipfile import ZipFile, ZipInfo

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover - environment diagnostic
    raise SystemExit("python-docx is required. Use the bundled workspace Python or install python-docx.") from exc


TEMPLATES = {
    "srs": "01-software-requirements-specification-template.docx",
    "sds": "02-software-solution-design-template.docx",
    "tas": "03-software-test-and-acceptance-template.docx",
}
ACRONYMS = {"srs": "SRS", "sds": "SDS", "tas": "TAS"}
TOTAL_DXA = 9026


def replace_text(paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for extra in paragraph.runs[1:]:
            extra._element.getparent().remove(extra._element)
    else:
        paragraph.add_run(value)


def retain_frontmatter(doc, title: str, args) -> None:
    body = doc._element.body
    children = [child for child in list(body) if child.tag != qn("w:sectPr")]
    if len(children) < 10:
        raise RuntimeError("DOCX template does not contain the expected cover and TOC frontmatter")
    for child in children[10:]:
        body.remove(child)

    paragraphs = doc.paragraphs
    replace_text(paragraphs[0], args.project_name)
    replace_text(paragraphs[1], title)
    replace_text(paragraphs[2], ACRONYMS[args.kind])

    if args.kind == "sds":
        cover_rows = [
            ("项目／产品名称", args.project_name),
            ("业务领域／功能模块", args.domain or args.module or "不适用"),
            ("文档编号", args.document_number),
            ("文档版本／状态", f"{args.version}／{args.status}"),
            ("编制／批准日期", f"{args.date}／{args.approval}"),
        ]
    else:
        cover_rows = [
            ("项目／产品名称", args.project_name),
            ("文档编号", args.document_number),
            ("文档版本", args.version),
            ("文档状态", args.status),
            ("编制／批准日期", f"{args.date}／{args.approval}"),
        ]
    cover = doc.tables[0]
    if len(cover.rows) != 5 or len(cover.columns) != 2:
        raise RuntimeError("DOCX template cover table must be 5x2")
    for row, values in zip(cover.rows, cover_rows):
        for cell, value in zip(row.cells, values):
            replace_text(cell.paragraphs[0], value)

    note = next((p for p in doc.paragraphs if "〈填写〉" in p.text or "示例行" in p.text), None)
    if note:
        replace_text(note, "本文已按标准模板填写；执行、批准和证据信息在正式使用时补充。")


def split_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def separator(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def set_repeat_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    if trpr.find(qn("w:tblHeader")) is None:
        node = OxmlElement("w:tblHeader")
        node.set(qn("w:val"), "true")
        trpr.append(node)


def widths(rows: list[list[str]], total: int = TOTAL_DXA) -> list[int]:
    count = max(len(row) for row in rows)
    weights = []
    for column in range(count):
        values = [row[column] if column < len(row) else "" for row in rows]
        weights.append(max(4, min(42, max((len(re.sub(r"[`*]", "", value)) for value in values), default=4))))
    minimum = max(620, min(1100, total // max(count * 3, 1)))
    remaining = max(0, total - minimum * count)
    result = [minimum + int(remaining * weight / sum(weights)) for weight in weights]
    result[-1] += total - sum(result)
    return result


def format_table(table, column_widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    tblpr = table._tbl.tblPr
    for tag, value in (("w:tblW", sum(column_widths)), ("w:tblInd", 100)):
        node = tblpr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tblpr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in column_widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    set_repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            width = column_widths[min(column_index, len(column_widths) - 1)]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "D9EAF7")
                cell._tc.get_or_add_tcPr().append(shading)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                if row_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)


def add_inline(paragraph, text: str) -> None:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    token = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    cursor = 0
    for match in token.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        value = match.group(0)
        run = paragraph.add_run(value[1:-1] if value.startswith("`") else value[2:-2])
        if value.startswith("`"):
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_markdown(doc, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    skipped_title = False
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("# ") and not skipped_title:
            skipped_title = True
            index += 1
            continue
        if line.startswith("```"):
            code = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            index += 1
            paragraph = doc.add_paragraph()
            for line_index, value in enumerate(code):
                if line_index:
                    paragraph.add_run().add_break()
                run = paragraph.add_run(value)
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
            continue
        if line.startswith("|") and index + 1 < len(lines):
            first, second = split_row(line), split_row(lines[index + 1])
            if separator(second):
                rows = [first]
                index += 2
                while index < len(lines) and lines[index].strip().startswith("|"):
                    rows.append(split_row(lines[index]))
                    index += 1
                table = doc.add_table(rows=0, cols=max(len(row) for row in rows))
                for source in rows:
                    cells = table.add_row().cells
                    for cell_index, cell in enumerate(cells):
                        value = source[cell_index] if cell_index < len(source) else ""
                        cell.text = ""
                        add_inline(cell.paragraphs[0], value)
                format_table(table, widths(rows))
                doc.add_paragraph()
                continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            paragraph = doc.add_paragraph(style=f"Heading {len(heading.group(1)) - 1}")
            add_inline(paragraph, heading.group(2))
            index += 1
            continue
        if line.startswith(">"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.22)
            paragraph.paragraph_format.space_after = Pt(6)
            add_inline(paragraph, line[1:].strip())
            index += 1
            continue
        bullet = re.match(r"^\s*-\s+(.+)$", raw)
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", raw)
        if bullet or numbered:
            paragraph = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            add_inline(paragraph, (bullet or numbered).group(1))
            index += 1
            continue
        paragraph = doc.add_paragraph()
        add_inline(paragraph, line)
        index += 1


def copy_info(info: ZipInfo) -> ZipInfo:
    clone = ZipInfo(info.filename, info.date_time)
    clone.compress_type = info.compress_type
    clone.comment = info.comment
    clone.extra = info.extra
    clone.internal_attr = info.internal_attr
    clone.external_attr = info.external_attr
    clone.create_system = info.create_system
    clone.flag_bits = info.flag_bits
    return clone


def assemble(template: Path, working: Path, output: Path) -> list[str]:
    with ZipFile(working) as source:
        replacements = {
            name: source.read(name)
            for name in ("word/document.xml", "word/header1.xml")
            if name in source.namelist()
        }
    temporary = output.with_suffix(".docx.tmp")
    with ZipFile(template) as original, ZipFile(temporary, "w") as result:
        for info in original.infolist():
            result.writestr(copy_info(info), replacements.get(info.filename, original.read(info.filename)))
    os.replace(temporary, output)
    return sorted(replacements)


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--kind", choices=TEMPLATES, required=True)
    parser.add_argument("--markdown", type=Path, required=True)
    parser.add_argument("--template", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--project-name", required=True)
    parser.add_argument("--module")
    parser.add_argument("--domain")
    parser.add_argument("--document-number", required=True)
    parser.add_argument("--version", default="V0.1")
    parser.add_argument("--status", default="草稿")
    parser.add_argument("--date", required=True)
    parser.add_argument("--approval", default="待批准")
    args = parser.parse_args()

    template = args.template or Path(__file__).resolve().parent.parent / "assets" / "templates" / TEMPLATES[args.kind]
    markdown = args.markdown.read_text(encoding="utf-8")
    title = next((line[2:].strip() for line in markdown.splitlines() if line.startswith("# ")), None)
    if not title:
        raise SystemExit("Markdown must contain one H1 document title")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="software-spec-docx-") as temporary:
        working = Path(temporary) / "working.docx"
        shutil.copy2(template, working)
        doc = Document(working)
        retain_frontmatter(doc, title, args)
        if doc.sections and doc.sections[0].header.paragraphs:
            replace_text(doc.sections[0].header.paragraphs[0], f"{args.project_name}  |  {ACRONYMS[args.kind]}")
        add_markdown(doc, markdown)
        doc.save(working)
        changed = assemble(template, working, args.output)
    print(
        f"created={args.output.resolve()} sha256={digest(args.output)} "
        f"template_parts_replaced={','.join(changed)}"
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
